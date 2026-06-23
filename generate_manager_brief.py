from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = "Contract_AI_Manager_Technical_Brief.docx"


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, color="D9DEE7", size="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
      borders = OxmlElement("w:tcBorders")
      tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_text(cell, text, bold=False, color="222222", size=8.3):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Aptos"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_border(cell)


def add_title(doc, text, subtitle):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    r.bold = True
    r.font.name = "Aptos Display"
    r.font.size = Pt(22)
    r.font.color.rgb = RGBColor(46, 46, 56)
    p.paragraph_format.space_after = Pt(2)

    p2 = doc.add_paragraph()
    r2 = p2.add_run(subtitle)
    r2.font.name = "Aptos"
    r2.font.size = Pt(9.5)
    r2.font.color.rgb = RGBColor(91, 99, 112)
    p2.paragraph_format.space_after = Pt(8)


def add_heading(doc, text, color="2E2E38"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    r.bold = True
    r.font.name = "Aptos Display"
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor.from_string(color)


def add_body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.03
    r = p.add_run(text)
    r.font.name = "Aptos"
    r.font.size = Pt(8.6)
    r.font.color.rgb = RGBColor(48, 55, 65)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style=None)
        p.paragraph_format.left_indent = Inches(0.14)
        p.paragraph_format.first_line_indent = Inches(-0.14)
        p.paragraph_format.space_after = Pt(1.5)
        p.paragraph_format.line_spacing = 1.0
        r = p.add_run("• ")
        r.font.name = "Aptos"
        r.font.size = Pt(8.2)
        r.font.color.rgb = RGBColor(255, 230, 0)
        rr = p.add_run(item)
        rr.font.name = "Aptos"
        rr.font.size = Pt(8.2)
        rr.font.color.rgb = RGBColor(48, 55, 65)


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    cell = table.cell(0, 0)
    shade_cell(cell, "FFF7B0")
    set_cell_border(cell, "E6D200", "8")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(title)
    r.bold = True
    r.font.name = "Aptos"
    r.font.size = Pt(8.7)
    r.font.color.rgb = RGBColor(46, 46, 56)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    rr = p2.add_run(body)
    rr.font.name = "Aptos"
    rr.font.size = Pt(8.1)
    rr.font.color.rgb = RGBColor(48, 55, 65)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_table(doc, headers, rows, widths=None, font_size=7.6):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, bold=True, color="FFFFFF", size=7.7)
        shade_cell(hdr[i], "2E2E38")
        if widths:
            hdr[i].width = widths[i]
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            set_cell_text(cells[i], val, size=font_size)
            if i == 0:
                shade_cell(cells[i], "F7F7F8")
            if widths:
                cells[i].width = widths[i]
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_pipeline(doc):
    headers = ["Layer", "Implementation", "Code location"]
    rows = [
        ["React AI workspace", "Single-page Contract Assistant with upload, compare, report prompts, chat, stats rail.", "gtsm-frontend/src/components/ContractUpload.jsx; App.css"],
        ["FastAPI API layer", "Routes for upload, compare, chat/RAG, analytics, history, workflow logs.", "app/main.py; app/routes/*.py"],
        ["Agentic workflow", "Adaptive orchestrator coordinates ingestion, clause analysis, risk, insight, department, memory, communication.", "app/agents/adaptive_orchestrator.py"],
        ["RAG memory", "Contract clauses embedded and stored in FAISS, then retrieved for chatbot answers.", "app/services/vector_store.py; app/routes/chat.py"],
        ["ML/LLM layer", "SentenceTransformer embeddings + local Ollama Llama 3.2 inference through OpenAI-compatible client.", "app/services/embedding_service.py; app/services/llm_service.py"],
    ]
    add_table(doc, headers, rows, [Inches(1.2), Inches(3.15), Inches(2.4)], font_size=7.2)


def build_doc():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.42)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)

    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(8.6)

    # Header/footer
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hr = header.add_run("Contract AI System | Manager Technical Brief")
    hr.font.name = "Aptos"
    hr.font.size = Pt(7)
    hr.font.color.rgb = RGBColor(91, 99, 112)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer.add_run("Prepared for technical review | 23 June 2026")
    fr.font.name = "Aptos"
    fr.font.size = Pt(7)
    fr.font.color.rgb = RGBColor(91, 99, 112)

    add_title(
        doc,
        "Contract AI System: Agentic AI Contract Intelligence Platform",
        "Three-page technical briefing for manager demo: what has been built, where it lives in code, and why it matters."
    )
    add_callout(
        doc,
        "Elevator pitch",
        "The system is now a single contract-first AI workspace: upload one PDF for Q&A, upload two PDFs for semantic comparison, retrieve clauses through RAG, generate risk/insight outputs, and view executive dashboard statistics from the same interface."
    )
    add_heading(doc, "1. What has been built so far")
    add_bullets(doc, [
        "Frontend has been consolidated into a polished single-page Contract AI workspace after login; code entry is App.jsx, main UI in ContractUpload.jsx, styling in App.css.",
        "Contract assistant supports one-contract processing, two-contract comparison, suggested prompts, detailed-report generation prompts, reset/new analysis, user header, and dashboard-statistics navigation.",
        "Backend exposes FastAPI routes for upload, compare, chat/RAG, analytics, dashboard history, workflow logs, departments, audit, escalation, and memory.",
        "Email delivery has been intentionally deactivated: email_service.py is now a no-send stub and compare/agent flows no longer show email UI messaging."
    ])
    add_heading(doc, "2. Architecture at a glance")
    add_pipeline(doc)

    doc.add_page_break()

    add_title(
        doc,
        "Agentic AI Architecture and ML Components",
        "The strongest technical story: modular agents + RAG + semantic similarity + local LLM inference."
    )
    add_heading(doc, "3. Agent workflow")
    add_table(
        doc,
        ["Agent / component", "Responsibility", "Implementation file"],
        [
            ["Adaptive Orchestrator", "Selects workflow agents, clears/logs workflow steps, persists memory/history, returns selected_agents and workflow_output.", "app/agents/adaptive_orchestrator.py"],
            ["Ingestion Agent", "Extracts PDF text, cleans content, splits document into clauses.", "app/agents/ingestion_agent.py; app/services/pdf_services.py"],
            ["Clause Agent", "Classifies clauses with LLM, generates embeddings, adds clauses into vector store.", "app/agents/clause_agent.py; app/services/clause_service.py"],
            ["Risk Agent", "Prompts LLM to return strict JSON risk_level, severity, impact, action, explanation.", "app/agents/risk_agent.py"],
            ["Insight Agent", "Turns risk + department context into executive summary, impact, recommendation, negotiation advice, next steps.", "app/agents/insight_agent.py"],
            ["Memory / history", "Stores clauses, risks, workflow logs, contract history, executive metrics.", "app/memory/*; app/services/history_service.py; app/services/workflow_tracker.py"],
        ],
        [Inches(1.45), Inches(3.5), Inches(2.0)],
        font_size=7.1
    )
    add_heading(doc, "4. ML / AI stack")
    add_table(
        doc,
        ["Capability", "Model / method", "Where implemented"],
        [
            ["Local LLM inference", "Ollama OpenAI-compatible endpoint with configurable OLLAMA_MODEL; default llama3.2:3b; temperature 0.2.", "app/services/llm_service.py"],
            ["Embeddings", "SentenceTransformer all-MiniLM-L6-v2, cached with lru_cache, 384-dimensional embeddings.", "app/services/embedding_service.py"],
            ["Vector retrieval", "FAISS IndexFlatL2 over stored clauses; search_clauses retrieves top-k relevant clauses for chat.", "app/services/vector_store.py"],
            ["RAG chatbot", "Chat route retrieves relevant clauses and injects them into a legal-assistant prompt.", "app/routes/chat.py"],
            ["Semantic comparison", "Cosine similarity between clause embeddings; risk bands: High < .70, Medium < .85, otherwise Low.", "app/services/comparison_service.py"],
            ["Report generation", "ReportLab PDF output for contract analysis results.", "app/services/report_service.py"],
        ],
        [Inches(1.45), Inches(3.45), Inches(2.0)],
        font_size=7.0
    )

    doc.add_page_break()

    add_title(
        doc,
        "Demo Narrative, Code Map, and Manager Talking Points",
        "Use this page as the speaking track for tomorrow's walkthrough."
    )
    add_heading(doc, "5. User-facing demo flow")
    add_bullets(doc, [
        "Login -> lands directly on Contract AI workspace, not a generic dashboard. This positions contract analysis as the core product.",
        "New analysis -> resets files, chat, and processed state. Dashboard statistics -> opens detailed charts/KPIs already built in ExecutiveCommandCenter.",
        "Process Contract -> POST /upload, stores clauses into FAISS for chatbot Q&A. Compare Contracts -> POST /compare, embeds both contracts and generates comparison report metadata.",
        "Chatbot -> POST /chat, uses retrieved clauses + workspace result to answer contract-specific questions and report prompts."
    ])
    add_heading(doc, "6. Feature-to-code map")
    add_table(
        doc,
        ["Feature", "Primary code location", "Technical note"],
        [
            ["One-page React workspace", "gtsm-frontend/src/App.jsx; components/ContractUpload.jsx", "App now routes directly to the contract workspace after login."],
            ["Professional UI system", "gtsm-frontend/src/App.css", "ChatGPT-style layout, left rail, user header, integrated stats, assistant composer."],
            ["Dashboard stats", "components/ExecutiveCommandCenter.jsx; ExecutiveKPIs.jsx; app/services/analytics_service.py", "KPI cards, charts, history, escalation feed, agent metrics."],
            ["Upload + RAG indexing", "app/routes/uploads.py; app/services/vector_store.py", "PDF clauses embedded once and inserted into FAISS."],
            ["Comparison engine", "app/routes/compare.py; app/services/comparison_service.py", "Embedding similarity and risk banding for clause mismatch analysis."],
            ["LLM risk/insight agents", "app/agents/risk_agent.py; insight_agent.py", "Structured JSON risk extraction and executive insight generation."],
        ],
        [Inches(1.55), Inches(2.45), Inches(2.95)],
        font_size=7.0
    )
    add_heading(doc, "7. Strong technical talking points")
    add_bullets(doc, [
        "Architecture is modular: UI, API routes, agents, services, memory, and analytics are separated so each layer can be improved independently.",
        "RAG is implemented end-to-end: clauses are embedded, stored in FAISS, retrieved by semantic search, then passed into the LLM prompt for grounded responses.",
        "Comparison is semantic, not string-only: clause embeddings are compared via cosine similarity, which handles wording changes better than exact matching.",
        "The platform is enterprise-friendly: local Ollama inference supports private contract data, audit/workflow tracking exists, and email was safely disabled for demo control.",
        "Next upgrade path is clear: persistent vector DB, native DOCX/XLSX parsing, role-based access, async job queue for faster uploads, and stronger evaluation metrics."
    ])

    doc.save(OUT)


if __name__ == "__main__":
    build_doc()
