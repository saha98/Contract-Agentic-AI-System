import sys
from datetime import datetime

sys.path.append(r"venv\Lib\site-packages")

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_DOCX = "Backend_Agentic_Workflow_Deep_Dive_2026-06-24.docx"
OUT_TXT = "Backend_Agentic_Workflow_Deep_Dive_2026-06-24.txt"


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, color="D9DEE7", size="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)

    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_text(cell, text, bold=False, color="222222", size=8.4):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Aptos"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_border(cell)


def add_title(doc, title, subtitle):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(title)
    run.bold = True
    run.font.name = "Aptos Display"
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(29, 37, 58)

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(8)
    run2 = p2.add_run(subtitle)
    run2.font.name = "Aptos"
    run2.font.size = Pt(9.5)
    run2.font.color.rgb = RGBColor(88, 96, 110)


def add_heading(doc, text, level=1):
    sizes = {
        1: 14,
        2: 11.5,
        3: 10.2,
    }
    colors = {
        1: "1D253A",
        2: "2E3A52",
        3: "4C566A",
    }
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5 if level == 1 else 3)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Aptos Display"
    run.font.size = Pt(sizes.get(level, 10))
    run.font.color.rgb = RGBColor.from_string(colors.get(level, "2E2E38"))


def add_body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.03
    run = p.add_run(text)
    run.font.name = "Aptos"
    run.font.size = Pt(8.7)
    run.font.color.rgb = RGBColor(48, 55, 65)


def add_bullets(doc, items, bullet_color="E6D200"):
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.14)
        p.paragraph_format.first_line_indent = Inches(-0.14)
        p.paragraph_format.space_after = Pt(1.2)
        p.paragraph_format.line_spacing = 1.0
        run = p.add_run("• ")
        run.font.name = "Aptos"
        run.font.size = Pt(8.2)
        run.font.color.rgb = RGBColor.from_string(bullet_color)
        body = p.add_run(item)
        body.font.name = "Aptos"
        body.font.size = Pt(8.2)
        body.font.color.rgb = RGBColor(48, 55, 65)


def add_callout(doc, title, body, fill="FFF8C9", border="D9C300"):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    set_cell_border(cell, border, "8")

    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(title)
    run.bold = True
    run.font.name = "Aptos Display"
    run.font.size = Pt(8.8)
    run.font.color.rgb = RGBColor(29, 37, 58)

    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.0
    run2 = p2.add_run(body)
    run2.font.name = "Aptos"
    run2.font.size = Pt(8.2)
    run2.font.color.rgb = RGBColor(48, 55, 65)

    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_table(doc, headers, rows, widths=None, font_size=7.6):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        set_cell_text(header_cells[index], header, bold=True, color="FFFFFF", size=7.7)
        shade_cell(header_cells[index], "1D253A")
        if widths:
            header_cells[index].width = widths[index]

    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            set_cell_text(cells[index], value, size=font_size)
            if index == 0:
                shade_cell(cells[index], "F6F7F9")
            if widths:
                cells[index].width = widths[index]

    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_page_break(doc):
    doc.add_page_break()


def build_route_flow_table():
    return [
        [
            "Contracts tab - single contract quick review",
            "POST /upload",
            "ContractUpload.jsx line 258",
            "uploads.py -> pdf_services.py -> clause_service.py -> services/vector_store.py",
            "Lightweight clause extraction, category preview, category breakdown, FAISS context for chat",
        ],
        [
            "Adaptive workflow tab - comprehensive contract intelligence",
            "POST /adaptive-workflow",
            "AdaptiveWorkflowPage.jsx line 161",
            "adaptive.py -> agents/adaptive_orchestrator.py -> agents + models + memory + history",
            "Full workflow_output with ML prediction, SHAP, governance, compliance, top risks, executive summary",
        ],
        [
            "Contracts tab - two contract comparison",
            "POST /compare",
            "ContractUpload.jsx line 296",
            "compare.py -> comparison_service.py -> insight_service.py -> report_service.py",
            "Clause similarity comparison, risk banding, generated PDF report path and filename",
        ],
        [
            "Download comparison report",
            "GET /compare-report",
            "ContractUpload.jsx line 403",
            "compare.py -> FileResponse",
            "Streams the comparison PDF saved under data/",
        ],
        [
            "Follow-up question after upload/compare",
            "POST /chat",
            "ContractUpload.jsx line 343 / 367",
            "chat.py -> services/vector_store.py -> llm_service.py",
            "Semantic retrieval over uploaded clauses plus LLM answer",
        ],
        [
            "Conversation memory mode",
            "POST /conversation",
            "Optional / not current main UI",
            "conversation.py -> agents/conversation_agent.py -> session_memory + FAISS + LLM",
            "Multi-turn chat with session-specific memory",
        ],
    ]


def build_agent_rows():
    return [
        [
            "adaptive_orchestrator.py",
            "make_json_safe() line 65; adaptive_orchestrator() line 111",
            "Live enterprise orchestrator for /adaptive-workflow. Clears logs, selects agents, runs ingestion -> memory -> clause -> features -> consensus -> ML -> XAI -> department -> risk -> governance -> compliance -> insight -> executive -> optional communication, then saves contract history.",
            "Called only by app/routes/adaptive.py after file upload.",
        ],
        [
            "ingestion_agent.py",
            "ingestion_agent() line 14",
            "Extracts raw PDF text, cleans it, and splits it into clauses using app/services/pdf_services.py.",
            "Used by adaptive_orchestrator.py, orchestrator.py, graphs/nodes.py, langgraph/nodes.py.",
        ],
        [
            "clause_agent.py",
            "clause_agent() line 13",
            "Filters trivial clauses, classifies each clause with classify_clause_llm(), generates embeddings, and inserts clauses into the FAISS clause store.",
            "Used by adaptive_orchestrator.py and graph prototypes.",
        ],
        [
            "risk_agent.py",
            "risk_agent() line 22",
            "Per clause, retrieves legal policy context from Chroma RAG store, builds a JSON-only risk prompt, calls the LLM, parses risk_level / severity / impact / action / explanation, and applies a fallback on failure.",
            "Used by adaptive_orchestrator.py and graph prototypes.",
        ],
        [
            "insight_agent.py",
            "insight_agent() line 20",
            "Aggregates all risk results with department hints into one LLM prompt that requests executive summary, top risks, negotiation advice, governance recommendations, and next steps.",
            "Used by adaptive_orchestrator.py and older orchestrator.py.",
        ],
        [
            "department_agent.py",
            "department_agent() line 16",
            "Keyword-routes clauses to Finance, Legal, or Risk based on payment, confidentiality/data, termination, and liability terms.",
            "Called during adaptive workflow before insight generation.",
        ],
        [
            "governance_agent.py",
            "governance_agent() line 4",
            "Converts risk_score into governance_score, decides whether escalation is required, and assigns review_level Standard vs Executive.",
            "Called during adaptive workflow and graph workflow.",
        ],
        [
            "compliance_agent.py",
            "compliance_agent() line 4",
            "Runs keyword detection across full processed clause text to flag GDPR, HIPAA, and SOX relevance.",
            "Called during adaptive workflow and graph workflow.",
        ],
        [
            "executive_agent.py",
            "executive_agent() line 10",
            "Calls the LLM to produce an executive summary from the ML risk payload, governance result, and compliance result.",
            "Called during adaptive workflow and graph workflow.",
        ],
        [
            "communication_agent.py",
            "communication_agent() line 11",
            "Wraps report_service.generate_report() and returns report path plus an email-disabled status object.",
            "Only added when the adaptive query contains email, notify, or communication.",
        ],
        [
            "conversation_agent.py",
            "conversation_agent() line 17",
            "Retrieves prior session memory, searches the FAISS clause store, builds a conversational prompt, calls the LLM, and saves the exchange.",
            "Called by /conversation, not by the main adaptive route.",
        ],
        [
            "feedback_agent.py",
            "feedback_agent() line 13; retrieve_feedback_history() line 30",
            "Stores tool feedback in memory and exposes retrieval for admin-only history.",
            "Called by /feedback and /feedback/history.",
        ],
        [
            "escalation_agent.py",
            "escalation_agent() line 11",
            "Looks up a company department in SQLite, creates a WorkflowLog row, and returns assignment details. Payment clauses map to Finance Team; all others default to Legal Team.",
            "Called only by /escalate.",
        ],
        [
            "orchestrator.py",
            "orchestrate_contract_workflow() line 14",
            "Older simple sequential workflow: ingestion -> clause -> risk -> insight -> communication.",
            "Not wired to any current FastAPI route.",
        ],
        [
            "tool_calling_agent.py",
            "tool_calling_agent() line 12",
            "Prototype agent that asks the LLM to choose among registered tools such as search_clauses, generate_report, send_email, and llm_reasoning.",
            "Not wired to main UI; used by test_agent.py.",
        ],
        [
            "test_agent.py",
            "Standalone test harness",
            "Manual script that calls tool_calling_agent() with a sample question.",
            "Used for local experimentation only.",
        ],
    ]


def build_graph_rows():
    return [
        [
            "contract_state.py",
            "ContractState TypedDict line 4",
            "Full graph state schema containing file_path, clauses, processed_clauses, contract_features, ML output, explainability, risks, governance, compliance, insights, and executive summary.",
            "Supports app/graphs/contract_workflow.py.",
        ],
        [
            "nodes.py",
            "ingestion_node() line 28 through executive_node() line 150",
            "Thin wrappers that adapt existing agents and models into graph nodes.",
            "Used by app/graphs/contract_workflow.py.",
        ],
        [
            "contract_workflow.py",
            "workflow definition line 12; graph compile line 130",
            "Builds a StateGraph with the full pipeline Ingestion -> Clause -> Feature -> Prediction -> XAI -> Risk -> Governance -> Compliance -> Department -> Insight -> Executive.",
            "Prototype / architecture artifact, not the active API path.",
        ],
    ]


def build_langgraph_rows():
    return [
        [
            "contract_state.py",
            "ContractState TypedDict line 4",
            "Smaller state object focused on user_query, file_path, clauses, processed_clauses, risks, departments, compliance, governance, insights, executive_summary, and communication.",
            "Used by app/langgraph/workflow.py.",
        ],
        [
            "nodes.py",
            "ingestion_node() line 22; clause_node() line 33; risk_node() line 48",
            "Reduced graph node set used for an earlier langgraph prototype.",
            "Used by app/langgraph/workflow.py.",
        ],
        [
            "workflow.py",
            "build_graph() line 21",
            "Creates a three-step graph: ingestion -> clause -> risk -> END.",
            "Invoked by graph_runner.py.",
        ],
        [
            "router.py",
            "route_workflow() line 1",
            "Keyword-based router that returns compliance, executive, or risk based on the user_query.",
            "Interesting design intent but not currently used by the active workflow.",
        ],
        [
            "graph_runner.py",
            "run_graph() line 6",
            "Compiles the langgraph workflow and invokes it with user_query and file_path.",
            "Used by test_langgraph.py, not by FastAPI routes.",
        ],
    ]


def build_kb_rows():
    return [
        ["payment_policy.txt", "Enterprise guidance for 15-45 days low risk, 60-90 moderate, >90 escalate to Finance."],
        ["payment_terms.txt", "Short finance guidance on enterprise-preferred 30-day cycle and escalation after 90 days."],
        ["liability_policy.txt", "Rules for unlimited liability, capped liability, and legal escalation triggers."],
        ["termination_policy.txt", "Balanced termination, cure periods, and unilateral termination risk guidance."],
        ["confidentiality_policy.txt", "Confidential information, privacy obligations, and compliance risk guidance."],
        ["gdpr_policy.txt", "Personal data, DPA references, cross-border transfers, breach notification, and data subject rights."],
        ["cybersecurity_policy.txt", "Security controls, encryption, incident notice, and certification expectations."],
        ["audit_rights_policy.txt", "Scope, frequency, confidentiality, notice period, and audit cost allocation guidance."],
        ["idemnity_policy.txt", "Mutual vs one-sided indemnity and capping expectations."],
        ["intellectual_property_policy.txt", "Ownership, reuse rights, source code ownership, and IP dispute exposure."],
    ]


def build_memory_rows():
    return [
        [
            "contract_memory.py",
            "save_contract() line 3; get_contracts() line 10",
            "Stores uploaded contracts and clause arrays in an in-memory list named contract_store.",
            "Adaptive workflow uses it through save_contract() and memory_agent().",
        ],
        [
            "risk_memory.py",
            "save_risks() line 3; get_risks() line 7",
            "Stores risk outputs in an in-memory list named risk_store.",
            "Adaptive workflow writes here after risk_agent() completes.",
        ],
        [
            "feedback_memory.py",
            "save_feedback() line 4; get_feedback_history() line 21",
            "Stores tool feedback in an in-memory list.",
            "Used by feedback_agent.py and /feedback routes.",
        ],
        [
            "session_memory.py",
            "save_conversation() line 4; get_conversation_history() line 20",
            "Stores chat exchanges in a session_id keyed dictionary.",
            "Used by conversation_agent.py.",
        ],
        [
            "memory_agent.py",
            "memory_agent() line 13",
            "Returns a consolidated snapshot of contracts, risks, feedback, and memory_status.",
            "Called inside adaptive_orchestrator.py.",
        ],
    ]


def build_model_rows():
    return [
        [
            "contract_feature_extractor.py",
            "extract_payment_days() line 4; extract_features() line 50",
            "Transforms processed clauses into a structured feature vector for ML, including payment_days plus 15 binary legal/compliance features.",
            "Called by adaptive_orchestrator.py and app/graphs/nodes.py.",
        ],
        [
            "xgboost_risk_model.py",
            "train_model() line 17; load_model() line 76; predict_risk() line 89",
            "Primary ML inference model using XGBClassifier. Trains from contract_training_dataset.csv if the pickle is absent, otherwise loads models/xgboost_risk_model.pkl.",
            "Called by adaptive_orchestrator.py, app/graphs/nodes.py, model_consensus.py, and explainability.py.",
        ],
        [
            "random_forest_risk_model.py",
            "train_model() line 17; predict_risk() line 70",
            "Secondary ensemble model used only for model consensus output.",
            "Called by model_consensus.py.",
        ],
        [
            "logistic_risk_model.py",
            "train_model() line 17; predict_risk() line 65",
            "Baseline linear model used only for model consensus output.",
            "Called by model_consensus.py.",
        ],
        [
            "model_consensus.py",
            "get_consensus_prediction() line 14",
            "Runs XGBoost, RandomForest, and LogisticRegression side by side and returns all three outputs plus a simple predictions list.",
            "Called by adaptive_orchestrator.py.",
        ],
        [
            "explainability.py",
            "explain_prediction() line 9",
            "Uses shap.TreeExplainer on the XGBoost model and returns top_factors intended to explain which features influenced the prediction.",
            "Called by adaptive_orchestrator.py and app/graphs/nodes.py.",
        ],
        [
            "risk_scoring_model.py",
            "extract_risk_features() line 4; score_contract() line 130",
            "Rule-based contract scoring layer that converts clause-level risk levels into overall risk_score, escalation_probability, and contract_health.",
            "Called by adaptive_orchestrator.py and app/graphs/nodes.py.",
        ],
        [
            "generate_training_data.py",
            "generate_dataset() line 5",
            "Synthetic training data generator for the ML models. Creates contract_training_dataset.csv with feature columns plus risk_class.",
            "Offline training support only; not called by API routes.",
        ],
    ]


def build_rag_rows():
    return [
        [
            "embeddings.py",
            "get_embedding_model() line 6; generate_embedding() line 19",
            "Loads SentenceTransformer all-MiniLM-L6-v2 for the knowledge base RAG flow and returns embeddings as Python lists.",
            "Used by ingest_documents.py and retriever.py.",
        ],
        [
            "vector_store.py",
            "PersistentClient setup line 3; get_collection() line 12",
            "Creates a persistent ChromaDB collection at vector_db/contract_knowledge_base.",
            "Used by KB ingestion and RAG retrieval.",
        ],
        [
            "ingest_documents.py",
            "ingest_document() line 13; ingest_folder() line 46",
            "Indexes knowledge base text files into Chroma with embeddings and source metadata.",
            "Typically run manually through load_kb.py.",
        ],
        [
            "retriever.py",
            "retrieve_context() line 10",
            "Embeds a query clause, performs a Chroma similarity search, logs the retrieved material, and returns concatenated context text.",
            "Used directly by risk_agent.py.",
        ],
    ]


def build_route_rows():
    return [
        [
            "uploads.py",
            "POST /upload at line 18",
            "Writes the PDF to data/uploads, extracts clauses, embeds them, LLM-classifies categories, stores them in FAISS, and returns total_clauses + sample_output + category_breakdown. This is the lightweight single-contract path, not the full enterprise workflow.",
            "Triggered from Contracts tab when only one file is uploaded.",
        ],
        [
            "compare.py",
            "POST /compare at line 64; GET /compare-report at line 72; compare_files() at line 99",
            "Processes two PDFs, embeds clauses, computes cosine similarity, generates risk insights, creates a comparison PDF, and exposes a secure report download route.",
            "Triggered from Contracts tab when two files are uploaded.",
        ],
        [
            "adaptive.py",
            "POST /adaptive-workflow at line 18",
            "Saves the uploaded contract to data/ and forwards the query plus file path into adaptive_orchestrator().",
            "Triggered from Adaptive Workflow page.",
        ],
        [
            "chat.py",
            "POST /chat at line 16",
            "Semantic clause retrieval over the FAISS store plus LLM answer generation.",
            "Triggered by follow-up questions in the Contracts tab.",
        ],
        [
            "conversation.py",
            "POST /conversation at line 24",
            "Calls conversation_agent() for multi-turn session-based chat.",
            "Available API path, not current main UI path.",
        ],
        [
            "feedback.py",
            "POST /feedback at line 34; GET /feedback/history at line 45",
            "Stores tool feedback and restricts history retrieval to admin role.",
            "Utility and admin route.",
        ],
        [
            "workflow_logs.py",
            "GET /workflow-logs at line 7",
            "Returns the in-memory workflow tracker logs.",
            "Used for monitoring / debugging.",
        ],
        [
            "memory.py",
            "GET /memory at line 19",
            "Returns in-memory contract, risk, and feedback stores.",
            "Monitoring / inspection path.",
        ],
        [
            "analytics.py",
            "GET /analytics at line 12; GET /executive-metrics at line 22",
            "Returns per-agent runtime analytics and executive KPI aggregates.",
            "Dashboard support route.",
        ],
        [
            "history.py",
            "GET /contract-history at line 11",
            "Reads persistent contract history rows from SQLite and formats them for the frontend.",
            "Review history and executive dashboard support route.",
        ],
        [
            "auth.py",
            "POST /signup at line 64; POST /login at line 94",
            "User creation, password hashing, password verification, and JWT issuance.",
            "Authentication route layer.",
        ],
        [
            "departments.py",
            "POST /departments at line 33; GET /departments/{company} at line 62",
            "Creates and fetches company-specific department contact records in SQLite.",
            "Supports escalation routing.",
        ],
        [
            "escalation.py",
            "POST /escalate at line 17",
            "Calls escalation_agent() to create workflow assignments.",
            "Manual escalation API path.",
        ],
        [
            "audit_logs.py",
            "GET /audit-logs at line 23",
            "Returns WorkflowLog ORM rows from SQLite.",
            "Audit support route.",
        ],
        [
            "agent_monitor.py",
            "GET /agent-monitor at line 6",
            "Returns a static hard-coded status payload for several agents.",
            "Presentation / placeholder monitoring route.",
        ],
        [
            "escalation_dashboard.py",
            "GET /escalation-dashboard at line 6",
            "Returns static hard-coded summary and escalation sample rows.",
            "Presentation / placeholder dashboard route.",
        ],
        [
            "audit_dashboard.py",
            "GET /audit-dashboard at line 6",
            "Returns static hard-coded audit events.",
            "Presentation / placeholder dashboard route.",
        ],
        [
            "department_dashboard.py",
            "GET /department-dashboard at line 6",
            "Returns static hard-coded department workload objects.",
            "Presentation / placeholder dashboard route.",
        ],
        [
            "executive_dashboard.py",
            "POST /executive-dashboard/store at line 9; GET /executive-dashboard at line 19",
            "Stores dashboard records in an in-memory list and returns them.",
            "Dashboard support route with non-persistent storage.",
        ],
        [
            "__init__.py",
            "Package init",
            "Minimal package initializer that imports uploads.router. app/main.py imports route modules directly, so this file is not the real runtime router map.",
            "Package support only.",
        ],
    ]


def build_service_rows():
    return [
        [
            "pdf_services.py",
            "extract_text_from_pdf() line 5; clean_text() line 15; split_into_clauses() line 26",
            "Foundational PDF parsing and clause splitting layer used by upload, compare, and ingestion_agent.",
            "Core preprocessing service.",
        ],
        [
            "clause_service.py",
            "classify_clause() line 12; classify_clause_llm() line 28; get_model() line 62; generate_embedding() line 72",
            "Provides heuristic and LLM clause categorization plus a SentenceTransformer embedding model for clause processing.",
            "Used by uploads.py, compare.py, and clause_agent.py.",
        ],
        [
            "embedding_service.py",
            "get_embedding_model() line 5; generate_embedding() line 11",
            "Lightweight cached SentenceTransformer embedding service for the FAISS clause store path.",
            "Used by services/vector_store.py and clause_agent.py.",
        ],
        [
            "vector_store.py",
            "add_clause() line 18; add_clause_with_embedding() line 25; search_clauses() line 34",
            "In-memory FAISS IndexFlatL2 clause store for uploaded contract content and semantic chat retrieval.",
            "Used by upload/compare/clause/chat/conversation/tool paths.",
        ],
        [
            "llm_service.py",
            "ask_llm() line 16",
            "Central LLM gateway using the OpenAI SDK against a local Ollama-compatible endpoint. Default model is llama3.2:3b with temperature 0 and max_tokens 1500.",
            "Used by clause classification, risk, insights, executive summary, chat, conversation, and tool calling.",
        ],
        [
            "comparison_service.py",
            "compare_clauses() line 3",
            "Computes cosine similarity between clause embeddings and converts similarity into High / Medium / Low comparison risk bands.",
            "Used by compare.py.",
        ],
        [
            "insight_service.py",
            "generate_insights() line 1",
            "Converts comparison scores into end-user-friendly reason and recommendation objects.",
            "Used by compare.py.",
        ],
        [
            "report_service.py",
            "generate_report() line 4",
            "Creates PDF reports with ReportLab under the data/ directory.",
            "Used by communication_agent.py and compare.py.",
        ],
        [
            "history_service.py",
            "save_contract_history() line 10; get_contract_history() line 45",
            "Persists contract history in SQLite and returns records for dashboards/history pages.",
            "Used by adaptive_orchestrator.py and history/analytics routes.",
        ],
        [
            "workflow_tracker.py",
            "log_step() line 8; get_logs() line 40; clear_logs() line 45",
            "In-memory execution tracking with timestamps and per-agent runtime calculation.",
            "Used by adaptive_orchestrator.py, risk_agent.py, governance_agent.py, compliance_agent.py, insight_agent.py, executive_agent.py, and workflow log analytics.",
        ],
        [
            "analytics_service.py",
            "calculate_agent_metrics() line 9; calculate_executive_metrics() line 61",
            "Builds aggregate runtime metrics from workflow_log and executive KPIs from contract_history.",
            "Used by analytics.py.",
        ],
        [
            "risk_history_service.py",
            "save_risk_history() line 5; get_risk_history() line 37",
            "SQLite persistence helper for clause-level risk history.",
            "Currently defined but not wired into the live adaptive path.",
        ],
        [
            "email_service.py",
            "send_email_report() line 1",
            "Intentional no-send stub that returns a disabled status payload.",
            "Used by tool registry and available for future reactivation.",
        ],
        [
            "__init__.py",
            "Package init",
            "No business logic; only package support.",
            "Support file.",
        ],
    ]


def build_ml_stack_rows():
    return [
        [
            "Local LLM inference",
            "Ollama via OpenAI-compatible SDK",
            "app/services/llm_service.py line 16",
            "All text generation and reasoning: clause categorization fallback, risk JSON, insight report, executive summary, chat, conversation, tool selection.",
        ],
        [
            "Clause embeddings for contract workspace",
            "SentenceTransformer all-MiniLM-L6-v2",
            "app/services/embedding_service.py line 5; app/services/clause_service.py line 62",
            "Embeds uploaded clauses for FAISS similarity search and comparison.",
        ],
        [
            "Knowledge base embeddings",
            "SentenceTransformer all-MiniLM-L6-v2",
            "app/rag/embeddings.py line 6",
            "Embeds static policy text for RAG retrieval inside risk_agent.",
        ],
        [
            "Vector search over uploaded clauses",
            "FAISS IndexFlatL2",
            "app/services/vector_store.py line 12 onward",
            "Grounds /chat and /conversation with previously uploaded clause context.",
        ],
        [
            "Persistent knowledge base vector store",
            "ChromaDB PersistentClient",
            "app/rag/vector_store.py line 3",
            "Holds legal policy documents used by the risk agent as reference material.",
        ],
        [
            "RAG clause risk grounding",
            "Embedding retrieval + prompt augmentation",
            "app/rag/retriever.py line 10; app/agents/risk_agent.py line 43",
            "Retrieves policy context before each risk prompt so the LLM reasons against enterprise legal guidance.",
        ],
        [
            "Feature engineering",
            "Regex + keyword-derived binary features",
            "app/models/contract_feature_extractor.py line 50",
            "Transforms processed clauses into structured inputs for XGBoost, RandomForest, and LogisticRegression.",
        ],
        [
            "Primary ML classifier",
            "XGBoost XGBClassifier",
            "app/models/xgboost_risk_model.py line 17 and line 89",
            "Produces risk_class, risk_score, and escalation_probability for the adaptive workflow.",
        ],
        [
            "Secondary ensemble models",
            "RandomForestClassifier and LogisticRegression",
            "app/models/random_forest_risk_model.py line 17; app/models/logistic_risk_model.py line 17",
            "Used to create a model consensus view for the adaptive dashboard.",
        ],
        [
            "Explainable AI",
            "SHAP TreeExplainer",
            "app/models/explainability.py line 19",
            "Generates top_factors intended to explain which features influenced the XGBoost result.",
        ],
        [
            "Rule-based risk scoring",
            "Weighted risk aggregation",
            "app/models/risk_scoring_model.py line 130",
            "Turns clause-level risk labels into overall contract risk metrics and health score.",
        ],
        [
            "Graph orchestration concept",
            "LangGraph StateGraph",
            "app/graphs/contract_workflow.py line 12; app/langgraph/workflow.py line 21",
            "Represents the same agent pipeline as an explicit state graph, though not currently bound to live API routes.",
        ],
    ]


def build_talking_points():
    return [
        "If your manager asks what the main runtime path is, answer: app/main.py mounts the API routers, and the real full-report path is /adaptive-workflow -> adaptive.py -> adaptive_orchestrator.py -> agents + models + memory + history.",
        "If he asks what happens after a normal contract upload in the Contracts tab, answer: /upload only does lightweight clause extraction, LLM/heuristic categorization, embeddings, and preview output. The comprehensive manager-facing report comes from /adaptive-workflow, not /upload.",
        "If he asks what makes the system agentic, answer: specialized agents own separate responsibilities, share workflow state, call different tools/models, and produce chained outputs instead of relying on one monolithic prompt. Then add the honest caveat that most of the orchestration order is fixed and only communication-agent selection is dynamic today.",
        "If he asks where ML is really used, answer: feature extraction in contract_feature_extractor.py feeds XGBoost, RandomForest, and LogisticRegression, while SHAP in explainability.py explains the XGBoost output and rule-based scoring in risk_scoring_model.py creates governance-ready metrics.",
        "If he asks where RAG is used, answer: the uploaded clause store uses FAISS for /chat, while the legal policy knowledge base uses Chroma and app/rag/retriever.py to ground risk_agent clause analysis.",
    ]


def build_text_report():
    sections = [
        "Backend Agentic Workflow Deep Dive",
        "Prepared for meeting on 24 June 2026.",
        "",
        "Quick position:",
        "The live enterprise report path is the adaptive workflow, not the simple upload route. /adaptive-workflow saves the contract, runs the adaptive orchestrator, extracts clauses, engineers features, performs ML risk prediction, runs SHAP explainability, generates rule-based risk metrics, checks governance and compliance, produces an executive summary, and persists contract history.",
        "",
        "Most important honest technical caveats:",
        "- The system is agentic in architecture, but only partially dynamic in routing today.",
        "- FAISS clause memory and several dashboard payloads are in-memory and reset on restart.",
        "- The legal policy knowledge base must be manually ingested through load_kb.py before RAG context is available.",
        "- The XGBoost risk_score currently uses predict_proba()[0][1], which is the probability for class 1 rather than the predicted class probability in a four-class model.",
        "- The SHAP explainability implementation currently flattens the multiclass SHAP output incorrectly and therefore under-reports feature impacts for many features.",
        "",
        "Use the DOCX for the complete file-by-file breakdown.",
    ]
    return "\n".join(sections)


def build_doc():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.42)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(8.7)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_run = header.add_run("Contract AI System | Backend Meeting Deep Dive")
    header_run.font.name = "Aptos"
    header_run.font.size = Pt(7)
    header_run.font.color.rgb = RGBColor(90, 98, 110)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run("Prepared on 24 June 2026 for technical manager walkthrough")
    footer_run.font.name = "Aptos"
    footer_run.font.size = Pt(7)
    footer_run.font.color.rgb = RGBColor(90, 98, 110)

    add_title(
        doc,
        "Adaptive Workflow Backend Deep Dive",
        "Detailed meeting brief covering exact trigger paths, backend folder responsibilities, ML/XAI components, and why the platform is agentic AI.",
    )

    add_callout(
        doc,
        "Most important high-level answer",
        "In the current product, the simple Contracts tab upload route is a lightweight ingest-and-preview path. The full enterprise contract intelligence report is produced by the Adaptive Workflow route, which calls adaptive_orchestrator() and then combines ingestion, clause intelligence, RAG-grounded risk analysis, feature engineering, XGBoost prediction, SHAP explainability, rule-based governance metrics, compliance checks, department mapping, LLM insight generation, and contract-history persistence.",
    )

    add_heading(doc, "1. Executive Overview", level=1)
    add_bullets(
        doc,
        [
            "Application boot starts in app/main.py. The FastAPI app is created at line 14 and routers are mounted at lines 27-45.",
            "The current frontend behavior matters: ContractUpload.jsx uses /upload for one-file review and /compare for two-file comparison, while AdaptiveWorkflowPage.jsx uses /adaptive-workflow for the comprehensive manager-facing dashboard.",
            "The system combines two retrieval stacks: FAISS over uploaded clauses for chat, and Chroma over policy documents for risk-agent RAG grounding.",
            "Persistent business history lives in SQLite through history_service.py and database models, while several memory and dashboard objects are still in-memory runtime structures.",
        ],
    )

    add_heading(doc, "2. Latest End-to-End Flow As The Product Works Today", level=1)
    add_table(
        doc,
        ["User action", "Backend endpoint", "Frontend trigger", "Main backend chain", "What the user gets"],
        build_route_flow_table(),
        [Inches(1.42), Inches(1.05), Inches(1.15), Inches(2.55), Inches(1.95)],
        font_size=7.05,
    )

    add_heading(doc, "2.1 Single-contract quick review in the Contracts tab", level=2)
    add_bullets(
        doc,
        [
            "ContractUpload.jsx posts the PDF to POST /upload at line 258 in the frontend.",
            "app/routes/uploads.py line 19 receives the file, saves it under data/uploads, and immediately calls extract_text_from_pdf(), clean_text(), and split_into_clauses().",
            "For every clause, uploads.py generates an embedding, asks classify_clause_llm() for a category, stores the clause plus embedding into the FAISS store through add_clause_with_embedding(), and builds category_breakdown.",
            "The route returns filename, total_clauses, sample_output, and category_breakdown. This is intentionally lightweight and does not call the adaptive orchestrator, XGBoost, SHAP, governance, compliance, or executive summary logic.",
            "The practical effect is: after /upload, the chat route can answer follow-up questions because the uploaded clauses now exist in the FAISS semantic store.",
        ],
    )

    add_heading(doc, "2.2 Comprehensive enterprise workflow for one uploaded contract", level=2)
    add_bullets(
        doc,
        [
            "AdaptiveWorkflowPage.jsx posts query + file to POST /adaptive-workflow.",
            "app/routes/adaptive.py writes the file to data/{filename} and forwards user_query plus file_path into adaptive_orchestrator().",
            "adaptive_orchestrator.py clears workflow logs, constructs workflow_data, and initializes selected_agents with clause_agent, risk_agent, and insight_agent. communication_agent is added only if the query contains email, notify, or communication.",
            "ingestion_agent() extracts, cleans, and splits the PDF into clauses, then save_contract() stores the raw clause set in in-memory contract memory.",
            "memory_agent() snapshots existing contracts, risks, and feedback into workflow_output['memory'].",
            "clause_agent() classifies each clause with the LLM, embeds it, and writes it into the FAISS clause index.",
            "extract_features() converts processed clause content into a structured feature vector; get_consensus_prediction() runs XGBoost, RandomForest, and LogisticRegression; predict_risk() generates the main ML result; explain_prediction() adds SHAP explainability.",
            "department_agent() assigns clauses to departments, risk_agent() performs RAG-grounded clause-by-clause risk analysis, and score_contract() converts the risk list into contract-level metrics.",
            "governance_agent(), compliance_agent(), insight_agent(), and executive_agent() transform the technical outputs into escalation, compliance, summary, and executive-facing report layers.",
            "Finally save_contract_history() writes a persistent SQLite history row and adaptive_orchestrator() returns selected_agents plus workflow_output.",
        ],
    )

    add_heading(doc, "2.3 Two-contract comparison flow", level=2)
    add_bullets(
        doc,
        [
            "When a second contract is present, ContractUpload.jsx posts both files to POST /compare.",
            "compare.py saves both files, uses process_file() to extract text, split clauses, generate embeddings, classify clauses heuristically, and write them into the FAISS store.",
            "comparison_service.compare_clauses() performs cosine similarity between every base clause and all candidate new clauses to find the best semantic match.",
            "Risk levels are derived from similarity thresholds: High for score < 0.7, Medium for score < 0.85, Low otherwise.",
            "insight_service.generate_insights() translates the technical score into user-facing reason and recommendation text.",
            "report_service.generate_report() creates a PDF under data/ and compare.py returns total_issues, risk_breakdown, report path, report filename, and top insight records.",
            "The frontend later hits GET /compare-report to stream the saved PDF back to the user.",
        ],
    )

    add_heading(doc, "2.4 Follow-up chat after upload or comparison", level=2)
    add_bullets(
        doc,
        [
            "ContractUpload.jsx sends a synthesized prompt to POST /chat.",
            "chat.py retrieves semantically similar clauses from the FAISS clause store through search_clauses().",
            "The retrieved context plus the user question are injected into a legal-assistant LLM prompt and answered by ask_llm().",
            "This means the chat capability is grounded in previously uploaded clauses, but it is separate from the adaptive orchestrator.",
        ],
    )

    add_heading(doc, "3. Exact Backend Trigger Map", level=1)
    add_body(doc, "The backend entrypoint is app/main.py. The FastAPI app is created at line 14. Routers are mounted in this order: uploads, compare, chat, feedback, conversation, adaptive, auth, departments, escalation, audit_logs, agent_monitor, escalation_dashboard, audit_dashboard, department_dashboard, workflow_logs, memory, analytics, executive_dashboard, and history.")
    add_callout(
        doc,
        "Important distinction for the meeting",
        "If someone asks 'what code is triggered when a contract is uploaded?', the accurate answer is: it depends on which page initiated the action. Normal upload in the Contracts tab triggers uploads.py only. Adaptive Workflow upload triggers adaptive.py and then the full adaptive_orchestrator pipeline.",
    )

    add_heading(doc, "4. Why This Is Agentic AI", level=1)
    add_bullets(
        doc,
        [
            "The workflow is decomposed into specialized agents instead of one giant prompt: ingestion, clause, risk, insight, department, governance, compliance, executive, memory, communication, escalation, conversation, feedback, and tool-calling agents all have distinct responsibilities.",
            "Agents share structured state. In the live path, adaptive_orchestrator.py builds workflow_data and passes outputs from one stage into downstream stages such as features -> ML -> XAI -> governance -> executive summary.",
            "Different agents call different tools and models: clause_agent uses embeddings and clause categorization, risk_agent uses RAG + LLM JSON extraction, insight_agent uses aggregate LLM reasoning, governance_agent uses deterministic rules, and executive_agent turns technical signals into management language.",
            "The system has memory and traceability. contract_memory, risk_memory, feedback_memory, session_memory, workflow_tracker, and history_service capture different kinds of state and auditability.",
            "There is also an explicit graph-based orchestration design in app/graphs and app/langgraph, which shows the architecture is intended as a stateful multi-step workflow and not a single request-response prompt wrapper.",
        ],
    )
    add_callout(
        doc,
        "Honest technical caveat",
        "This is best described as semi-agentic rather than fully autonomous. The architecture is clearly agent-based, but the live adaptive route still follows a mostly fixed sequence. Dynamic routing today mainly affects whether communication_agent is included based on the user query.",
    )

    add_heading(doc, "5. ML, AI, and Advanced Concepts Used", level=1)
    add_table(
        doc,
        ["Capability", "Model / method", "Code location", "How it is used in this system"],
        build_ml_stack_rows(),
        [Inches(1.32), Inches(1.42), Inches(1.75), Inches(3.1)],
        font_size=6.95,
    )

    add_heading(doc, "5.1 How XGBoost is trained and used", level=2)
    add_bullets(
        doc,
        [
            "Training data generation lives in app/models/generate_training_data.py. generate_dataset() synthesizes rows using payment_days and binary legal-feature flags such as liability_clause, termination_clause, confidentiality_clause, compliance_clause, audit_rights, indemnity_clause, cybersecurity_clause, gdpr_clause, insurance_clause, force_majeure, intellectual_property, subcontracting, vendor_risk, data_retention, and cross_border_transfer.",
            "app/models/xgboost_risk_model.py train_model() loads contract_training_dataset.csv, separates X from y, constructs XGBClassifier with n_estimators=300, max_depth=6, learning_rate=0.05, subsample=0.8, colsample_bytree=0.8, and random_state=42, then saves the fitted model to models/xgboost_risk_model.pkl.",
            "load_model() lazily loads the pickle and retrains only if the file does not exist.",
            "During live adaptive inference, extract_features() creates the current contract feature vector, predict_risk() converts it into a one-row DataFrame, runs model.predict() and model.predict_proba(), maps the class id to Low/Medium/High/Critical, and returns risk_class plus risk_score and escalation_probability.",
            "The adaptive dashboard shows this payload under workflow_output['ml_risk_prediction'].",
        ],
    )

    add_heading(doc, "5.2 How SHAP XAI is intended to work", level=2)
    add_bullets(
        doc,
        [
            "app/models/explainability.py load_model() reuses the XGBoost model and wraps it with shap.TreeExplainer.",
            "The current contract feature vector is converted into a one-row DataFrame, explainer.shap_values(df) is computed, and the code attempts to build a per-feature impact list.",
            "The intended product behavior is: identify the highest-impact features behind the ML risk output and surface them under workflow_output['explainability']['top_factors'].",
            "In the frontend, ExplainabilityPanel renders these top factors for the user as the XAI explanation layer.",
        ],
    )

    add_heading(doc, "5.3 How clause-level risk and prediction connect to contract reporting", level=2)
    add_bullets(
        doc,
        [
            "Clause-level semantic processing begins in clause_agent.py and app/services/clause_service.py. Clauses are categorized and embedded first.",
            "Clause-level risk generation happens in risk_agent.py, where each processed clause is combined with retrieved policy context and sent to the LLM for a structured JSON risk response.",
            "Contract-level ML prediction is separate: extract_features() compresses the full contract into structured signals, then XGBoost / RandomForest / LogisticRegression score that feature vector.",
            "Rule-based contract scoring in risk_scoring_model.py converts the clause-level risk list into a business-facing contract score and health score. Governance and executive summary then sit on top of those scores.",
            "That layered design is what produces both detailed clause insight and high-level executive output in the same workflow.",
        ],
    )

    add_page_break(doc)

    add_heading(doc, "6. Folder-by-Folder Backend Walkthrough", level=1)

    add_heading(doc, "6.1 app/agents", level=2)
    add_table(
        doc,
        ["File", "Key functions / lines", "Responsibility", "How it is triggered"],
        build_agent_rows(),
        [Inches(1.2), Inches(1.35), Inches(3.1), Inches(1.95)],
        font_size=6.9,
    )

    add_heading(doc, "6.2 app/graphs", level=2)
    add_table(
        doc,
        ["File", "Key functions / lines", "Responsibility", "How it is triggered"],
        build_graph_rows(),
        [Inches(1.2), Inches(1.4), Inches(3.0), Inches(2.0)],
        font_size=7.0,
    )
    add_body(doc, "Manager explanation: app/graphs is the richer graph-based architecture artifact. It mirrors the live enterprise workflow more closely than app/langgraph, but it is not currently called by any FastAPI route.")

    add_heading(doc, "6.3 app/langgraph", level=2)
    add_table(
        doc,
        ["File", "Key functions / lines", "Responsibility", "How it is triggered"],
        build_langgraph_rows(),
        [Inches(1.2), Inches(1.4), Inches(3.0), Inches(2.0)],
        font_size=7.0,
    )
    add_body(doc, "Manager explanation: app/langgraph is an earlier / lighter orchestration prototype. test_langgraph.py calls graph_runner.run_graph(), but the live API does not.")

    add_heading(doc, "6.4 app/knowledge_base", level=2)
    add_table(
        doc,
        ["File", "Purpose in the system"],
        build_kb_rows(),
        [Inches(2.0), Inches(5.8)],
        font_size=7.2,
    )
    add_body(doc, "These policy files are not queried directly by the frontend. They become usable only after manual ingestion into Chroma through load_kb.py, which calls app/rag/ingest_documents.py ingest_folder('app/knowledge_base/legal').")

    add_heading(doc, "6.5 app/memory", level=2)
    add_table(
        doc,
        ["File", "Key functions / lines", "Responsibility", "How it is triggered"],
        build_memory_rows(),
        [Inches(1.2), Inches(1.35), Inches(3.1), Inches(1.95)],
        font_size=7.0,
    )
    add_body(doc, "Important explanation: this folder is memory in the runtime sense, not long-term persistence. These stores reset when the backend restarts.")

    add_heading(doc, "6.6 app/models", level=2)
    add_table(
        doc,
        ["File", "Key functions / lines", "Responsibility", "How it is triggered"],
        build_model_rows(),
        [Inches(1.2), Inches(1.35), Inches(3.1), Inches(1.95)],
        font_size=7.0,
    )

    add_heading(doc, "6.7 app/rag", level=2)
    add_table(
        doc,
        ["File", "Key functions / lines", "Responsibility", "How it is triggered"],
        build_rag_rows(),
        [Inches(1.2), Inches(1.35), Inches(3.1), Inches(1.95)],
        font_size=7.0,
    )

    add_heading(doc, "6.8 app/routes", level=2)
    add_table(
        doc,
        ["File", "Key endpoint / lines", "Responsibility", "How it is triggered"],
        build_route_rows(),
        [Inches(1.1), Inches(1.4), Inches(3.25), Inches(1.85)],
        font_size=6.9,
    )

    add_heading(doc, "6.9 app/services", level=2)
    add_table(
        doc,
        ["File", "Key functions / lines", "Responsibility", "How it is triggered"],
        build_service_rows(),
        [Inches(1.1), Inches(1.45), Inches(3.2), Inches(1.85)],
        font_size=6.9,
    )

    add_page_break(doc)

    add_heading(doc, "7. Important Technical Caveats To Be Ready For", level=1)
    add_callout(
        doc,
        "Why this section matters",
        "Your manager is technical, so these are exactly the kinds of questions he may ask. Calling them out proactively usually strengthens credibility because it shows you understand both what is already working and what still needs tightening.",
        fill="FFE7CC",
        border="D08B00",
    )
    add_bullets(
        doc,
        [
            "The live adaptive workflow is agent-based, but selected_agents does not list every component that actually runs. governance, compliance, department, executive, and memory steps still execute even though the selected_agents array mainly shows clause, risk, insight, and sometimes communication.",
            "Several runtime stores are non-persistent: app/memory/* and app/services/vector_store.py are in-memory only, and app/routes/executive_dashboard.py also stores records in an in-memory list.",
            "The legal policy knowledge base is not auto-ingested on application startup. Someone must run load_kb.py for RAG policy retrieval to be populated.",
            "app/routes/adaptive.py imports require_role but does not currently apply Depends(require_role(...)), so the adaptive endpoint is not actually role-protected in its current form.",
            "app/routes/agent_monitor.py, escalation_dashboard.py, audit_dashboard.py, and department_dashboard.py return static hard-coded payloads, so those dashboards are presentation placeholders rather than live operational analytics.",
            "The comparison flow uses classify_clause() in compare.py, which is a heuristic classifier, while the upload/adaptive clause paths use classify_clause_llm(). That means clause categorization quality is not identical across all flows.",
            "department_agent.py only returns a department entry when keyword rules match. insight_agent.py aligns department_results to risk_results by index, so unmatched clauses fall back to department='General'.",
            "workflow_tracker.py is in-memory. Agent runtime analytics therefore represent the current server lifecycle, not a persistent historical performance warehouse.",
        ],
    )

    add_heading(doc, "7.1 ML/XAI caveats worth mentioning carefully", level=2)
    add_bullets(
        doc,
        [
            "The XGBoost model is a four-class classifier. In app/models/xgboost_risk_model.py, predict_risk() currently uses model.predict_proba(df)[0][1] as the probability basis for risk_score and escalation_probability. In a four-class setting, index 1 is only the class-1 probability, not necessarily the probability of the predicted class. So the current scalar risk_score should be described as a heuristic output rather than a mathematically perfect calibrated score.",
            "The SHAP code in app/models/explainability.py handles a multiclass SHAP tensor as if it were a single vector. On a real sample the SHAP output shape is (1, 16, 4), but the code repeatedly indexes [0] until only the first feature's class vector remains. This means top_factors currently surface meaningful values for only a subset of features and zero out many others. It is a valid explainability concept in the architecture, but the current implementation needs correction for production-grade XAI.",
            "generate_training_data.py creates synthetic rather than human-labeled legal outcomes. That is fine for demonstrating the architecture, but it should not be presented as a mature supervised model trained on a large reviewed contract corpus unless you explicitly say the training data is synthetic and rule-derived.",
            "There is also a label-generation bug risk in generate_training_data.py: risk_class is assigned only inside the cross_border_transfer branch, which means rows without that branch reuse the previous loop value instead of recomputing risk_class cleanly. The current saved dataset exists, but the generator logic should be fixed before any retraining story is presented as production-ready.",
        ],
    )

    add_heading(doc, "8. Suggested Speaking Track For The Meeting", level=1)
    add_bullets(doc, build_talking_points())

    add_heading(doc, "9. Clean Manager Summary", level=1)
    add_body(doc, "The backend is centered on FastAPI routes that branch into two main product behaviors: a lightweight contract workspace upload path and a richer adaptive workflow path. The richer path is where the enterprise demo value lives: it orchestrates multiple specialized agents, combines retrieval, embeddings, LLM reasoning, classical ML prediction, SHAP explainability, governance/compliance checks, and persistent history capture, and returns a structured workflow_output that the frontend renders into the dashboard.")
    add_body(doc, "From a technical-story standpoint, the strongest claims are: the architecture is modular, the platform uses both RAG and classical ML, the adaptive flow is traceable and stateful, and there is a clear separation between preprocessing, semantic retrieval, clause intelligence, risk analysis, governance, and executive reporting. The most important honest caveats are around persistence, dashboard placeholders, and the current ML/XAI calibration details.")

    doc.save(OUT_DOCX)


if __name__ == "__main__":
    build_doc()
    with open(OUT_TXT, "w", encoding="utf-8") as handle:
        handle.write(build_text_report())
    print(OUT_DOCX)
    print(OUT_TXT)
